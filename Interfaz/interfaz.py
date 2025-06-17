import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import tkinter.font as tkFont
from PIL import ImageTk, Image, ImageDraw
import os
import subprocess
from adbutils import adb
import cv2
import re
import time
import threading

import pandas as pd

EXCEL_PATH = "participantes.xlsx"


def inicializar_excel():
    if not os.path.exists(EXCEL_PATH):
        df = pd.DataFrame(columns=["nombre", "sesion", "escena"])
        df.to_excel(EXCEL_PATH, index=False)
    else:
        df = pd.read_excel(EXCEL_PATH)
        for col in ["nombre", "sesion", "escena"]:
            if col not in df.columns:
                df[col] = ""
        df.to_excel(EXCEL_PATH, index=False)


import subprocess
import shlex

archivo = "/sdcard/Movies/EscenariosNeutros/Montaña.jpg"
archivo_quoted = shlex.quote(archivo)

comando = f'adb shell rm {archivo_quoted}'
subprocess.run(comando, shell=True)



def limpiar_directorio_temporal():
    tmp_dir = "tmp"
    if os.path.exists(tmp_dir):
        for archivo in os.listdir(tmp_dir):
            ruta = os.path.join(tmp_dir, archivo)
            try:
                os.remove(ruta)
            except Exception as e:
                print(f"No se pudo borrar {ruta}: {e}")


def actualizar_ultima_escena(nombre, escena):
    df = pd.read_excel(EXCEL_PATH)
    idx = df[df["nombre"] == nombre].index
    if not idx.empty:
        df.at[idx[0], "escena"] = escena
        df.to_excel(EXCEL_PATH, index=False)


def agregar_participante_excel(nombre):
    df = pd.read_excel(EXCEL_PATH)
    if nombre not in df["nombre"].values:
        nuevo = pd.DataFrame([[nombre, 0]], columns=["nombre", "sesion"])
        df = pd.concat([df, nuevo], ignore_index=True)
        df.to_excel(EXCEL_PATH, index=False)


def obtener_sesion(nombre):
    df = pd.read_excel(EXCEL_PATH)
    fila = df[df["nombre"] == nombre]
    if not fila.empty:
        return int(fila["sesion"].values[0])
    return 0


def incrementar_sesion(nombre):
    df = pd.read_excel(EXCEL_PATH)
    idx = df[df["nombre"] == nombre].index
    if not idx.empty:
        df.at[idx[0], "sesion"] += 1
        df.to_excel(EXCEL_PATH, index=False)
        return df.at[idx[0], "sesion"]
    return None


# Carpeta para guardar pacientes
PARTICIPANTES_DIR = "participantes"

if not os.path.exists(PARTICIPANTES_DIR):
    os.makedirs(PARTICIPANTES_DIR)

# === Inicializacion ===
historial_botones = []
participantes = {}  # {nombre: sesión_actual}
participante_actual = None

videos_gafas = []
seconds = 0
running = False
imagenes_miniaturas = []
video_actual = None
video_inicio_segundos = None


def registrar_boton(nombre):
    tiempo = timer_var.get()
    historial_botones.append(f"[{tiempo}] {nombre}")

'''
def cargar_imagen():
    archivo = filedialog.askopenfilename(filetypes=[("Imágenes", "*.png;*.jpg;*.jpeg")])
    if archivo:
        img = Image.open(archivo)
        img = img.resize((250, 150))
        imagen = ImageTk.PhotoImage(img)
        imagen_label.config(image=imagen)
        imagen_label.image = imagen
'''

def generar_thumbnail(ruta_video):
    try:
        cap = cv2.VideoCapture(ruta_video)
        ret, frame = cap.read()
        cap.release()
        if ret:
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            img_pil = Image.fromarray(frame_rgb).resize((160, 90))
            return ImageTk.PhotoImage(img_pil)
    except Exception as e:
        print(f"Error generando thumbnail para {ruta_video}: {e}")
    return None

def abrir_video(ruta):
    global video_actual, video_inicio_segundos

    try:
        nombre_video = os.path.basename(ruta)

        # cerrar anterior
        if video_actual:
            duracion = seconds - video_inicio_segundos
            minutos = duracion // 60
            segundos_restantes = duracion % 60
            registrar_boton(f"Fin de video: {video_actual} (duración: {minutos:02}:{segundos_restantes:02})")

        # abrir nuevo
        subprocess.Popen([ruta], shell=True)
        registrar_boton(f"Inicio de video: {nombre_video}")

        video_actual = nombre_video
        video_inicio_segundos = seconds

        # actualizar excel
        if participante_actual and "ansioso" in ruta.lower():
            actualizar_ultima_escena(participante_actual, nombre_video)

    except Exception as e:
        print(f"Error al abrir el video: {e}")


def cargar_video():
    archivo = filedialog.askopenfilename(filetypes=[("Videos", "*.mp4;*.avi;*.mov;*.mkv")])
    if archivo:
        video_label.config(text=f"Video cargado:\n{archivo.split('/')[-1]}")
        videos_gafas.append(archivo)
        nombre = os.path.basename(archivo)
        btn = tk.Button(frame_videos_gafas, text=nombre, anchor="w", relief="groove",
                        command=lambda path=archivo: abrir_video(path))
        btn.pack(fill="x", pady=2, padx=2)


def mostrar_informe():
    ventana = tk.Toplevel(root)
    ventana.title("Informe de Sesión")
    ventana.geometry("400x400")

    tk.Label(ventana, text="Historial de botones pulsados:", font=("Segoe UI", 12, "bold")).pack(pady=10)

    frame_texto = tk.Frame(ventana)
    frame_texto.pack(fill="both", expand=True, padx=10, pady=(0, 5))

    scrollbar = tk.Scrollbar(frame_texto)
    scrollbar.pack(side="right", fill="y")

    texto = tk.Text(frame_texto, wrap="word", yscrollcommand=scrollbar.set, height=15)
    texto.pack(side="left", fill="both", expand=True)

    scrollbar.config(command=texto.yview)

    if historial_botones:
        for linea in historial_botones:
            texto.insert("end", f"• {linea}\n")
    else:
        texto.insert("end", "No se ha pulsado ningún botón aún.")

    texto.config(state="disabled")

    frame_boton = tk.Frame(ventana)
    frame_boton.pack(fill="x", pady=(0, 10))

    btn_exportar_pdf = tk.Button(frame_boton, text="Exportar a PDF", command=exportar_informe_pdf)
    btn_exportar_pdf.pack(side="left", padx=10)

    btn_exportar_word = tk.Button(frame_boton, text="Exportar a Word", command=exportar_informe_word)
    btn_exportar_word.pack(side="left", padx=10)


from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from datetime import datetime


def exportar_informe_pdf():
    if not historial_botones:
        messagebox.showinfo("Sin datos", "No hay historial para exportar.")
        return

    nombre_archivo = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf")],
        title="Guardar informe como PDF"
    )
    if not nombre_archivo:
        return

    try:
        c = canvas.Canvas(nombre_archivo, pagesize=A4)
        width, height = A4
        margen = 50
        y = height - margen
        nombre = entry_id.get()
        sesion = entry_sesion.get()
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margen, y, "Informe de sesión")
        y -= 30

        c.setFont("Helvetica", 12)
        c.drawString(margen, y, f"Paciente: {nombre}")
        y -= 20
        c.drawString(margen, y, f"Sesión Nº: {sesion}")
        y -= 30

        c.setFont("Helvetica", 12)

        fecha_actual = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        nombre = entry_id.get()
        sesion = entry_sesion.get()

        c.drawString(margen, y, f"Fecha: {fecha_actual}")
        y -= 20
        c.drawString(margen, y, f"Paciente: {nombre}")
        y -= 20
        c.drawString(margen, y, f"Sesión Nº: {sesion}")
        y -= 30
        c.drawString(margen, y, f"Fecha: {fecha_actual}")
        y -= 30

        c.setFont("Helvetica", 11)
        for linea in historial_botones:
            if y < margen + 30:
                c.showPage()
                y = height - margen
                c.setFont("Helvetica", 11)
            c.drawString(margen, y, f"- {linea}")
            y -= 20

        c.save()
        messagebox.showinfo("Informe exportado", "El informe se ha guardado correctamente.")
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo generar el PDF:\n{e}")


from docx import Document
from docx.shared import Inches
from datetime import datetime


def exportar_informe_word():
    if not historial_botones:
        messagebox.showinfo("Sin datos", "No hay historial para exportar.")
        return

    nombre_archivo = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx")],
        title="Guardar informe como Word"
    )
    if not nombre_archivo:
        return

    try:
        doc = Document()
        doc.add_heading('Informe de sesión', 0)

        fecha_actual = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        nombre = entry_id.get()
        sesion = entry_sesion.get()

        doc.add_paragraph(f"Fecha: {fecha_actual}")

        doc.add_heading('Historial de botones pulsados', level=1)
        for linea in historial_botones:
            doc.add_paragraph(f"- {linea}", style='List Bullet')

        doc.save(nombre_archivo)
        messagebox.showinfo("Informe exportado", "El informe se ha guardado correctamente en Word.")
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo generar el Word:\n{e}")


def reset_timer():
    global running, seconds
    running = False
    seconds = 0
    timer_var.set("00:00")


def actualizar_timer():
    global seconds
    if running:
        minutos = seconds // 60
        segundos = seconds % 60
        timer_var.set(f"{minutos:02}:{segundos:02}")
        seconds += 1
    root.after(1000, actualizar_timer)


def empezar_aplicacion():
    global running, seconds

    if participante_actual:
        nueva_sesion = incrementar_sesion(participante_actual)
        if nueva_sesion is not None:
            entry_sesion.configure(state="normal")
            entry_sesion.delete(0, tk.END)
            entry_sesion.insert(0, str(nueva_sesion))
            entry_sesion.configure(state="readonly")

    seconds = 0
    running = True


def parar_aplicacion():
    global running, video_actual, video_inicio_segundos
    running = False

    if video_actual:
        duracion = seconds - video_inicio_segundos
        minutos = duracion // 60
        segundos_restantes = duracion % 60
        registrar_boton(f"Fin de video: {video_actual} (duración: {minutos:02}:{segundos_restantes:02})")
        video_actual = None
        video_inicio_segundos = None


def parar_aplicacion():
    global running, video_actual, video_inicio_segundos
    running = False

    registrar_boton("PARAR APLICACIÓN")

    # Marca final del video #
    if video_actual:
        duracion = seconds - video_inicio_segundos
        minutos = duracion // 60
        segundos_restantes = duracion % 60
        registrar_boton(f"Fin de video: {video_actual} (duración: {minutos:02}:{segundos_restantes:02})")
        video_actual = None
        video_inicio_segundos = None

    # Registrar duración total de la sesión
    minutos = seconds // 60
    segundos_restantes = seconds % 60
    historial_botones.append(f"Duración total de la sesión: {minutos:02}:{segundos_restantes:02}")



# === Agregar participante

def agregar_participante():
    global participante_actual

    nuevo = tk.simpledialog.askstring("Nuevo participante", "Nombre del nuevo participante:")
    if nuevo:
        agregar_participante_excel(nuevo)
        combo_participantes['values'] = actualizar_lista_participantes()
        combo_participantes.set(nuevo)
        participante_actual = nuevo
        cargar_datos_participante(nuevo)


def cargar_datos_participante(nombre):
    global participante_actual
    participante_actual = nombre
    sesion = obtener_sesion(nombre)

    entry_id.delete(0, tk.END)
    entry_id.insert(0, participante_actual)

    entry_sesion.configure(state="normal")
    entry_sesion.delete(0, tk.END)
    entry_sesion.insert(0, str(sesion))
    entry_sesion.configure(state="readonly")



def crear_botones_desde_carpeta():
    carpeta = filedialog.askdirectory()
    if carpeta:
        archivos = [f for f in os.listdir(carpeta) if os.path.isfile(os.path.join(carpeta, f))]

        for widget in frame_archivos.winfo_children():
            widget.destroy()

        for archivo in archivos:
            btn = tk.Button(frame_archivos, text=archivo, anchor="w")
            btn.pack(fill="x", padx=5, pady=2)


def cargar_escenas_gafas():
    devices = [ip for ip in list_devices_tcpip() if ip.strip() != ""]
    device = adb.device(devices[0])
    output = device.shell('ls -1 /sdcard/Movies/AssetBundleScenes')
    assetbundles = []
    for f in output.split('\n'):
        f = f.strip()
        if f and not (f.endswith('.meta') or f.endswith('.manifest')):
            assetbundles.append(f)
    # print("Escenas:", assetbundles)
    return assetbundles


''' DANI
def crear_botones_videos_desde_carpeta():
    carpeta = filedialog.askdirectory()
    if carpeta:
        extensiones_validas = ('.mp4', '.avi', '.mov', '.mkv')
        archivos = [
            f for f in os.listdir(carpeta)
            if os.path.isfile(os.path.join(carpeta, f)) and f.lower().endswith(extensiones_validas)
        ]

        for widget in frame_galeria_videos.winfo_children():
            widget.destroy()

        imagenes_miniaturas.clear()
        columnas = 4
        fila = 0
        columna = 0

        for archivo in archivos:
            ruta_completa = os.path.join(carpeta, archivo)
            thumbnail = generar_thumbnail(ruta_completa)
            if thumbnail:
                btn = tk.Button(frame_galeria_videos, image=thumbnail,
                                command=lambda path=ruta_completa: abrir_video(path))
                btn.image = thumbnail
                btn.grid(row=fila, column=columna, padx=5, pady=5)
                columna += 1
                if columna >= columnas:
                    columna = 0
                    fila += 1

global frame_galeria_videos
'''

def generar_miniatura_local(ruta_video):
    try:
        cap = cv2.VideoCapture(ruta_video)
        ret, frame = cap.read()
        cap.release()
        if ret:
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            img_pil = Image.fromarray(frame_rgb).resize((160, 90))
            return img_pil
    except Exception as e:
        print(f"Error generando miniatura: {e}")
    return None

import shlex

def adb_push(local_path, remote_path):
    try:
        comando = f'adb push "{local_path}" "{remote_path}"'
        result = subprocess.run(
            shlex.split(comando),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        if result.returncode != 0:
            print(f"Error: {result.stderr}")
    except Exception as e:
        print(f"Error ejecutando adb push: {e}")

        if result.returncode != 0:
            print(f"Error: {result.stderr}")
    except Exception as e:
        print(f"Error ejecutando adb push: {e}")



import unicodedata

import shutil  # Para copiar archivos temporalmente

def normalizar_nombre(nombre):
    nombre = nombre.replace(" ", "_")
    nombre = nombre.replace("ñ", "n").replace("Ñ", "N")
    return nombre




def subir_videos_gafas(tipo="neutro"):
    if tipo == "ansioso":
        ruta_remota = "/sdcard/Movies/EscenariosAnsiosos"
    else:
        ruta_remota = "/sdcard/Movies/EscenariosNeutros"


    archivos = filedialog.askopenfilenames(
        title="Selecciona videos para subir a las gafas",
        filetypes=[("Archivos de video", "*.mp4 *.avi *.mov *.mkv")]
    )

    if not archivos:
        print("No se seleccionaron archivos.")
        return

    limpiar_directorio_temporal()



    ventana_progreso = tk.Toplevel()
    ventana_progreso.title("Subiendo videos")
    ventana_progreso.geometry("400x100")
    ventana_progreso.resizable(False, False)

    etiqueta_estado = tk.Label(ventana_progreso, text="Preparando...", anchor="w")
    etiqueta_estado.pack(fill="x", padx=10, pady=(10, 0))

    barra = ttk.Progressbar(ventana_progreso, maximum=len(archivos), mode="determinate")
    barra.pack(fill="x", padx=10, pady=10)

    def hilo_subida():

        for i, ruta_local in enumerate(archivos):
            nombre_original = os.path.basename(ruta_local)
            nombre_normalizado = normalizar_nombre(nombre_original)

            if nombre_normalizado != nombre_original:
                ruta_temp = os.path.join("tmp", nombre_normalizado)
                os.makedirs("tmp", exist_ok=True)
                shutil.copy(ruta_local, ruta_temp)
                ruta_local_usar = ruta_temp
            else:
                ruta_local_usar = ruta_local

            ruta_remota_video = f"{ruta_remota}/{nombre_normalizado}"

            def actualizar_estado(msg, progreso=None):
                def _actualizar():
                    etiqueta_estado.config(text=msg)
                    if progreso is not None:
                        barra['value'] = progreso
                ventana_progreso.after(0, _actualizar)

            try:
                actualizar_estado(f"Subiendo: {nombre_normalizado}", i)
                adb_push(ruta_local_usar, ruta_remota_video)

                miniatura = generar_miniatura_local(ruta_local)
                if miniatura:
                    nombre_thumb = os.path.splitext(nombre_normalizado)[0] + ".jpg"
                    ruta_thumb_local = os.path.join("tmp", nombre_thumb)
                    miniatura.save(ruta_thumb_local)
                    ruta_remota_thumb = f"{ruta_remota}/{nombre_thumb}"
                    adb_push(ruta_thumb_local, ruta_remota_thumb)

                actualizar_estado(f"Subido: {nombre_normalizado}", i + 1)

            except Exception as e:
                print(f"Error al subir {ruta_local}: {e}")
                actualizar_estado(f"Error con {nombre_normalizado}", i + 1)

        ventana_progreso.after(0, lambda: etiqueta_estado.config(text="Subida completada."))
        ventana_progreso.after(1000, ventana_progreso.destroy)

    threading.Thread(target=hilo_subida, daemon=True).start()


def cargar_videos_gafas():
    # print("hola")
    devices = [ip for ip in list_devices_tcpip() if ip.strip() != ""]
    device = adb.device(devices[0])
    output = device.shell('ls /sdcard/Movies/EscenariosNeutros')
    videos = []
    for f in output.split():
        f = f.strip()
        if f.endswith('.mp4'):
            videos.append(f)

    return videos

def archivo_existe_en_dispositivo(ruta):
    import subprocess
    try:
        resultado = subprocess.run(['adb', 'shell', 'ls', ruta],
                                   stdout=subprocess.PIPE,
                                   stderr=subprocess.PIPE,
                                   text=True)
        return resultado.returncode == 0
    except Exception as e:
        print(f"Error al comprobar existencia de {ruta}: {e}")
        return False

import tempfile

def copiar_imagen_desde_dispositivo(ruta_remota):
    try:
        tmp_dir = tempfile.gettempdir()
        nombre_archivo = os.path.basename(ruta_remota)
        ruta_local = os.path.join(tmp_dir, nombre_archivo)
        subprocess.run(['adb', 'pull', ruta_remota, ruta_local], check=True)
        return ruta_local
    except Exception as e:
        print(f"Error al copiar imagen desde dispositivo: {e}")
        return None

def crear_botones_videos_desde_carpeta():
    videos = cargar_videos_gafas()

    if not videos:
        print("No se encontraron videos. Saliendo de la función.")
        return

    for widget in frame_galeria_videos.winfo_children():
        widget.destroy()

    imagenes_miniaturas.clear()
    columnas = 4
    fila = 0
    columna = 0

    font_texto = tkFont.Font(family="Arial", size=10)

    img_transparente = Image.new('RGBA', (160, 90), (255, 255, 255, 0))
    img_transparente_tk = ImageTk.PhotoImage(img_transparente)
    imagenes_miniaturas.append(img_transparente_tk)

    for video in videos:
        base_name = video.rsplit('.', 1)[0]
        ruta_img_jpg = f"/sdcard/Movies/EscenariosNeutros/{base_name}.jpg"
        ruta_img_png = f"/sdcard/Movies/EscenariosNeutros/{base_name}.png"

        #print(f"Buscando imagen para video: {video}")

        ruta_imagen = None
        if archivo_existe_en_dispositivo(ruta_img_jpg):
            ruta_imagen = copiar_imagen_desde_dispositivo(ruta_img_jpg)
        elif archivo_existe_en_dispositivo(ruta_img_png):
            ruta_imagen = copiar_imagen_desde_dispositivo(ruta_img_png)
        #print(ruta_imagen)

        # Crear un subframe para miniatura + texto
        marco_video = tk.Frame(frame_galeria_videos)

        if ruta_imagen:
            try:
                img = Image.open(ruta_imagen)
                img.thumbnail((160, 90))
                thumbnail = ImageTk.PhotoImage(img)
                imagenes_miniaturas.append(thumbnail)

                btn = tk.Button(marco_video, image=thumbnail,
                                command=lambda path=f"/sdcard/Movies/EscenariosNeutros/{video}": abrir_video(path))
                btn.image = thumbnail
                btn.pack()

            except Exception as e:
                btn = tk.Button(marco_video, image=img_transparente_tk,
                                command=lambda path=f"/sdcard/Movies/EscenariosNeutros/{video}": abrir_video(path))
                btn.image = img_transparente_tk
                btn.pack()
        else:
            btn = tk.Button(marco_video, image=img_transparente_tk,
                            command=lambda path=f"/sdcard/Movies/EscenariosNeutros/{video}": abrir_video(path))
            btn.image = img_transparente_tk
            btn.pack()

        # Etiqueta con el nombre del video (sin extensión, más limpio)
        nombre_video_sin_extension = os.path.splitext(os.path.basename(video))[0]
        etiqueta = tk.Label(marco_video, text=nombre_video_sin_extension, font=font_texto, wraplength=160)
        etiqueta.pack()

        # Posicionar el marco en la cuadrícula
        marco_video.grid(row=fila, column=columna, padx=5, pady=5)

        columna += 1
        if columna >= columnas:
            columna = 0
            fila += 1

    #print("Finalizada la creación de botones.")

global frame_galeria_videos

# Extensiones a mostrar
EXTENSIONES_VISIBLES = ('.mp4', '.avi', '.mov', '.mkv')
# Extensiones a eliminar como "extra" junto con vídeos seleccionados
EXTENSIONES_EXTRA_BORRAR = ('.jpg', '.png')

def obtener_lista_videos():
    resultado = subprocess.run(
        ['adb', 'shell', 'ls', '/sdcard/Movies/EscenariosNeutros/'],
        stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
    )
    if resultado.returncode != 0:
        print("Error al obtener la lista de vídeos:", resultado.stderr)
        return []

    archivos = resultado.stdout.strip().splitlines()
    videos = [f for f in archivos if f.lower().endswith(EXTENSIONES_VISIBLES)]
    return videos


def borrar_videos_carpeta(videos_a_borrar):
    import shlex
    errores = []

    for video in videos_a_borrar:
        base = re.sub(r'\.[^.]+$', '', video)
        archivos_a_borrar = [f"{base}.mp4"] + [f"{base}{ext}" for ext in EXTENSIONES_EXTRA_BORRAR]

        for archivo in archivos_a_borrar:
            ruta = f"/sdcard/Movies/EscenariosNeutros/{archivo}"
            ruta_quoted = shlex.quote(ruta)

            # Ejecutar como string completo para evitar errores de codificación
            comando = f'adb shell rm {ruta_quoted}'

            try:
                resultado = subprocess.run(
                    comando,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    shell=True,
                    text=True,
                    encoding='utf-8'  # esto es importante
                )
                if resultado.returncode != 0 and "No such file" not in resultado.stderr:
                    errores.append(archivo)
                    print(f"Error al borrar {archivo}: {resultado.stderr}")
            except Exception as e:
                errores.append(archivo)
                print(f"Excepción al borrar {archivo}: {e}")

    if errores:
        messagebox.showerror("Error", f"No se pudieron borrar: {', '.join(errores)}")
    else:
        messagebox.showinfo("Éxito", "Vídeos borrados correctamente.")

def pestaña_videos_gafas():
    print("Crea pestaña con los videos en las gafas")
    videos = obtener_lista_videos()

    if not videos:
        messagebox.showwarning("Sin vídeos", "No se encontraron vídeos en las gafas.")
        return

    ventana = tk.Tk()
    ventana.title("Videos en Gafas")

    lista_videos = tk.Listbox(ventana, selectmode=tk.MULTIPLE, width=50)
    lista_videos.pack(padx=10, pady=10)

    for video in videos:
        lista_videos.insert(tk.END, video)

    def borrar_seleccionados():
        seleccion = [lista_videos.get(i) for i in lista_videos.curselection()]
        if not seleccion:
            messagebox.showwarning("Nada seleccionado", "Selecciona al menos un vídeo para borrar.")
            return
        confirmar = messagebox.askyesno("Confirmar", f"¿Seguro que quieres borrar {len(seleccion)} vídeo(s)?")
        if confirmar:
            borrar_videos_carpeta(seleccion)
            # Actualizar lista
            lista_videos.delete(0, tk.END)
            for video in obtener_lista_videos():
                lista_videos.insert(tk.END, video)
                crear_botones_videos_desde_carpeta()

    boton_borrar = tk.Button(ventana, text="Borrar seleccionados", command=borrar_seleccionados)
    boton_borrar.pack(pady=10)

    ventana.mainloop()

def crear_botones_imagenes_desde_carpeta():
    carpeta = filedialog.askdirectory()
    if carpeta:
        extensiones_validas = ('.png', '.jpg', '.jpeg', '.bmp', '.gif')
        archivos = [
            f for f in os.listdir(carpeta)
            if os.path.isfile(os.path.join(carpeta, f)) and f.lower().endswith(extensiones_validas)
        ]

        for widget in frame_imagenes_dinamicas.winfo_children():
            if isinstance(widget, tk.Button) and widget != btn_cargar_imagenes:
                widget.destroy()

        imagenes_miniaturas.clear()

        for archivo in archivos:
            ruta_completa = os.path.join(carpeta, archivo)
            try:
                img = Image.open(ruta_completa)
                img.thumbnail((60, 60))
                img_tk = ImageTk.PhotoImage(img)
                imagenes_miniaturas.append(img_tk)  # evitar que se borre de memoria

                btn = tk.Button(frame_imagenes_dinamicas, image=img_tk,
                                command=lambda path=ruta_completa: mostrar_imagen_en_panel(path))
                btn.pack(side="left", padx=3, pady=3)
            except Exception as e:
                print(f"Error cargando {archivo}: {e}")


##NUEVO MARTA##
def get_device_ip(device):
    try:
        device = adb.device(device)
        output = device.shell("ip route")
        match = re.search(r"(\d+\.\d+\.\d+\.\d+)$", output)
        return match.group(1) if match else None
    except Exception as e:
        '''
        print(f"Error getting IP: {e}")
        if "Can't find any android device/emulator" in e:
            messagebox.showwarning("Conecta las gafas por cable al ordenador")
        '''
        return None


def connect_devices_adb():
    devices = adb.device_list()
    items = [i.serial for i in devices if ":" in i.serial]  # Filtrar solo dispositivos TCP/IP
    try:
        device = adb.device(items)
        output = device.shell("ip route")
        match = re.search(r"(\d+\.\d+\.\d+\.\d+)$", output)
        device.tcpip(5555)
        time.sleep(2)
        for _ in range(1):
            subprocess.call(f"adb connect {match.group(1)}")
            return True
    except Exception as e:
        print(f"Error getting IP: {e}")
        if "Can't find any android device/emulator" in str(e):
            messagebox.showinfo("Gafas no conectadas", "Conecta las gafas por cable al ordenador y pulsa de nuevo 'Conectar dispositivo'.")
        return None


def conectar_y_actualizar():
    conectado = connect_devices_adb()
    if conectado:
        # Actualiza lista de dispositivos en el ComboBox
        nuevos_dispositivos = list_devices_tcpip()
        ip_combo['values'] = nuevos_dispositivos
        if nuevos_dispositivos:
            ip_combo.set(nuevos_dispositivos[0])
            start_battery_monitoring(nuevos_dispositivos[0], battery_entry, root)
        else:
            ip_combo.set("No devices found")
            battery_entry.configure(state="normal")
            battery_entry.delete(0, tk.END)
            battery_entry.insert(0, "No device found")
            battery_entry.configure(state="readonly")


def list_devices_tcpip():
    devices = adb.device_list()
    items = [i.serial for i in devices if ":" in i.serial]  # solo dispositivos TCP/IP
    return items


def start_scrcpy(serial, extra_args=None):
    try:
        scrcpy_path = r"C:\Users\danie\OneDrive\Escritorio\TFG-Final\scrcpy-win64-v3.2\scrcpy.exe"
        print("Starting scrcpy with device:", serial)
        command = [
            scrcpy_path,
            '-s', serial,
            '--crop', '1832:1920:0:0'
        ]
        if extra_args:
            command += extra_args
        print("Executing command:", command)

        def run_scrcpy():
            subprocess.run(command)

        scrcpy_thread = threading.Thread(target=run_scrcpy)
        scrcpy_thread.start()

    except Exception as e:
        print("Exception:", e)
        messagebox.showerror("Error", str(e))


def mirror_hmd_view():
    selected_indices = list_devices_tcpip()
    if selected_indices:
        start_scrcpy(selected_indices[0])
    else:
        print("No devices found")


def actualizar_bateria(device, battery_entry):
    device = adb.device(serial=device)
    if device:
        try:
            output = device.shell("dumpsys battery | grep level")
            battery_level = int(output.strip().split(":")[1])

            # Update Entry
            battery_entry.configure(state="normal")
            battery_entry.delete(0, tk.END)
            battery_entry.insert(0, f"{battery_level}%")

            # Update color based on battery level
            color = "red" if battery_level < 20 else "black"
            battery_entry.configure(fg=color)
            battery_entry.configure(state="readonly")

        except Exception as e:
            print(f"Error reading battery level: {e}")


def start_battery_monitoring(device, battery_entry, root, interval_ms=60000):
    def update():
        actualizar_bateria(device, battery_entry)
        root.after(interval_ms, update)

    update()


##TERMINA NUEVO MARTA##

def mostrar_imagen_en_panel(path):
    try:
        img = Image.open(path)
        img.thumbnail((400, 300))
        img_tk = ImageTk.PhotoImage(img)
        imagen_preview_label.config(image=img_tk)
        imagen_preview_label.image = img_tk
    except Exception as e:
        print(f"No se pudo mostrar la imagen: {e}")


def reorganizar_botones(event=None):
    ancho = root.winfo_width()
    if ancho < 800:
        for i, b in enumerate(botones):
            b.grid(row=i, column=0, sticky="ew", padx=5, pady=2)
        frame_botones.columnconfigure(0, weight=1)
    else:
        for i, b in enumerate(botones):
            b.grid(row=0, column=i, sticky="ew", padx=5, pady=2)
        for i in range(len(botones)):
            frame_botones.columnconfigure(i, weight=1)


'''
# === Setup device ===

devices = list_devices_tcpip()
if len(devices) > 0:
    ip = get_device_ip(devices)
    device = adb.device(serial=devices)
    alive = True
    client = scrcpy.Client(device=device, bitrate=1000000, max_fps=5)
    #client.add_listener(scrcpy.EVENT_INIT, self.on_init)
    #client.add_listener(scrcpy.EVENT_FRAME, self.on_frame)
    client.add_listener(scrcpy.EVENT_FRAME, on_frame2)

    #self.start_scrcpy_thread()  # Llamar al hilo de Scrcpy
else:
    print("No ADB devices found. Running interface without device.")
    device = None
    alive = True
'''
inicializar_excel()

# === Root window ===
root = tk.Tk()
root.title("Interfaz adaptable")


'''
root.attributes("-fullscreen", True)
'''
root.geometry("1200x700")

# === Barra herramientas ===
toolbar = tk.Frame(root, bd=1, relief="raised", bg="#f0f0f0")
toolbar.pack(side="top", fill="x")


def seleccionar_participante(event):
    seleccionado = combo_participantes.get()
    cargar_datos_participante(seleccionado)


def actualizar_lista_participantes():
    if not os.path.exists(EXCEL_PATH):
        return []
    df = pd.read_excel(EXCEL_PATH)
    return df["nombre"].tolist()

''''''
'''
btn_nuevo_participante = tk.Button(toolbar, text="Añadir participante", command=lambda: agregar_participante())
btn_nuevo_participante.pack(side="left", padx=2)
'''
'''
btn_abrir_img = tk.Button(toolbar, text="Abrir Imagen", command=cargar_imagen)
btn_abrir_img.pack(side="left", padx=2, pady=2)
'''
'''
btn_abrir_video = tk.Button(toolbar, text="Abrir Video", command=cargar_video)
btn_abrir_video.pack(side="left", padx=2, pady=2)
'''
'''
btn_guardar = tk.Button(toolbar, text="Guardar Sesión")
btn_guardar.pack(side="left", padx=2, pady=2)
'''
'''
btn_salir = tk.Button(toolbar, text="Salir", command=root.quit)
btn_salir.pack(side="right", padx=2, pady=2)
'''
'''
btn_ventana = tk.Button(toolbar, text="Modo ventana", command=lambda: root.attributes("-fullscreen", False))
btn_ventana.pack(side="left", padx=2, pady=2)
'''
# === Contenido principal ===
frame_contenido = tk.Frame(root)
frame_contenido.pack(fill="both", expand=True)

'''
frame_contenido.columnconfigure(0, weight=3, minsize=800)
frame_contenido.columnconfigure(1, weight=2, minsize=400)
'''
frame_contenido.columnconfigure(0, weight=1)

frame_contenido.rowconfigure(0, weight=1)

# Para quitar el scroll==
# === scrol izquierda ===
frame_scroll = tk.Frame(frame_contenido)
frame_scroll.grid(row=0, column=0, sticky="nsew")

frame_scroll.rowconfigure(0, weight=1)
frame_scroll.columnconfigure(0, weight=1)

canvas_izquierda = tk.Canvas(frame_scroll)
scrollbar_izquierda = tk.Scrollbar(frame_scroll, orient="vertical", command=canvas_izquierda.yview)
canvas_izquierda.configure(yscrollcommand=scrollbar_izquierda.set)

scrollbar_izquierda.pack(side="right", fill="y")
canvas_izquierda.pack(side="left", fill="both", expand=True)

# Marco real donde van los widgets
frame_izquierda = tk.Frame(canvas_izquierda)
canvas_izquierda.create_window((0, 0), window=frame_izquierda, anchor="nw", tags="contenedor")


# Ajustar scroll al tamaño del contenido
def ajustar_scroll(event):
    canvas_izquierda.configure(scrollregion=canvas_izquierda.bbox("all"))


def expandir_canvas(event):
    canvas_izquierda.itemconfig("contenedor", width=event.width)


frame_izquierda.bind("<Configure>", ajustar_scroll)
canvas_izquierda.bind("<Configure>", expandir_canvas)

frame_izquierda.columnconfigure(0, weight=1)

# === Frame superior: Datos paciente ===
frame_superior = tk.Frame(frame_izquierda)
frame_superior.grid(row=0, column=0, sticky="ew", padx=10, pady=5)
for i in range(4):
    frame_superior.columnconfigure(i, weight=1)

tk.Label(frame_superior, text="ID del paciente").grid(row=0, column=0, sticky="w")

tk.Label(frame_superior, text="Nº sesión").grid(row=0, column=2, sticky="w")

entry_id = tk.Entry(frame_superior)
entry_id.grid(row=0, column=1, sticky="ew", padx=5)

entry_sesion = tk.Entry(frame_superior, state="readonly")
entry_sesion.grid(row=0, column=3, sticky="ew", padx=5)

# === Botones principales + Timer ===
frame_botones = tk.Frame(frame_izquierda)
frame_botones.grid(row=1, column=0, sticky="ew", padx=10, pady=10)

botones = [
    tk.Button(frame_botones, text="EMPEZAR APLICACIÓN", bg="lightgreen"),
    tk.Button(frame_botones, text="PARAR APLICACIÓN", bg="red", fg="white")
]

frame_botones.columnconfigure(5, weight=1)

for b in botones:
    b.grid()

frame_timer = tk.LabelFrame(frame_botones, text="Timer")
frame_timer.grid(row=0, column=3, rowspan=2, padx=10)



# === Frame Participante ===
frame_participante = tk.LabelFrame(frame_botones, text="Participante")
frame_participante.grid(row=0, column=4, rowspan=2, padx=10, sticky="nsew")
frame_participante.columnconfigure(0, weight=1)

tk.Label(frame_participante, text="Selecciona:").grid(row=0, column=0, padx=5, pady=(5, 0), sticky="ew")

combo_participantes = ttk.Combobox(frame_participante, state="readonly")
combo_participantes.grid(row=1, column=0, padx=5, pady=2, sticky="ew")
combo_participantes.bind("<<ComboboxSelected>>", seleccionar_participante)

btn_nuevo_participante = tk.Button(frame_participante, text="Añadir", command=agregar_participante)
btn_nuevo_participante.grid(row=2, column=0, padx=5, pady=(2, 8), sticky="ew")


timer_var = tk.StringVar()
timer_var.set("00:00")
timer_label = tk.Label(frame_timer, textvariable=timer_var, font=("Courier", 16))
timer_label.pack(padx=5, pady=(5, 2))
btn_restart = tk.Button(frame_timer, text="Restart", command=reset_timer)
btn_restart.pack(pady=(0, 5))

botones[0].config(command=lambda: [registrar_boton("EMPEZAR APLICACIÓN"), empezar_aplicacion()])
botones[1].config(command=lambda: [registrar_boton("PARAR APLICACIÓN"), parar_aplicacion()])

'''
# === Frame tareas ===
frame_tareas = tk.Frame(frame_izquierda)
frame_tareas.grid(row=3, column=0, sticky="nsew", padx=10, pady=10)
frame_tareas.columnconfigure(1, weight=1)

tk.Label(frame_tareas, text="TAREA 1").grid(row=0, column=0, sticky="w", pady=5)
ttk.Combobox(frame_tareas, values=["Nivel 1", "Nivel 2", "Nivel 3"]).grid(row=0, column=1, sticky="ew", pady=5)

tk.Label(frame_tareas, text="TAREA 2").grid(row=1, column=0, sticky="w", pady=5)
ttk.Combobox(frame_tareas, values=["Nivel 1", "Nivel 2", "Nivel 3"]).grid(row=1, column=1, sticky="ew", pady=5)
'''

def cargar_videos_ansiosos():
    devices = [ip for ip in list_devices_tcpip() if ip.strip()]
    if not devices:
        return []
    device = adb.device(devices[0])
    output = device.shell('ls /sdcard/Movies/EscenariosAnsiosos')
    videos = []
    for f in output.split():
        f = f.strip()
        if f.endswith('.mp4'):
            videos.append(f)
    return videos


# === Archivos Gafas ===
frame_videos_gafas = tk.LabelFrame(frame_izquierda, text="Archivos de las Gafas", padx=5, pady=5)
frame_videos_gafas.grid(row=4, column=0, sticky="ew", padx=10, pady=(0, 10))
frame_videos_gafas.columnconfigure(0, weight=1)

# === Dispositivo ===
frame_dispositivo = tk.LabelFrame(frame_izquierda, text="Dispositivo", padx=5, pady=5)
frame_dispositivo.grid(row=5, column=0, sticky="ew", padx=10, pady=(0, 10))
frame_dispositivo.columnconfigure(1, weight=1)

tk.Label(frame_dispositivo, text="IP:").grid(row=0, column=0, sticky="e")
device_ips = [ip for ip in list_devices_tcpip() if ip.strip() != ""]
ip_combo = ttk.Combobox(frame_dispositivo, values=device_ips)
ip_combo.grid(row=0, column=1, sticky="ew", padx=5, pady=2)
if device_ips:
    ip_combo.current(0)
else:
    ip_combo.set("No devices found")

tk.Label(frame_dispositivo, text="Batería:").grid(row=1, column=0, sticky="e")
battery_entry = tk.Entry(frame_dispositivo, state="readonly", justify="center")
battery_entry.grid(row=1, column=1, sticky="ew", padx=5, pady=2)
battery_entry.insert(0, "42%")
if device_ips:
    start_battery_monitoring(device_ips[0], battery_entry, root)
else:
    battery_entry.insert(0, "No device found")

btn_devices = tk.Button(frame_dispositivo, text="Conectar dispositivo",
                        command=conectar_y_actualizar)  # connect_devices_adb()
btn_devices.grid(row=2, column=0, columnspan=2, pady=5, sticky="ew")

mirror_button = tk.Button(frame_dispositivo, text="Ver vista de las gafas", command=mirror_hmd_view)
mirror_button.grid(row=3, column=0, columnspan=2, pady=5, sticky="ew")

cargar_videos = tk.Button(frame_dispositivo, text="Acceder archivos gafas", command=cargar_escenas_gafas)
cargar_videos.grid(row=4, column=0, columnspan=2, pady=5, sticky="ew")

# === Mirror vista gafas ===

# === Volumen ===
frame_volumen = tk.LabelFrame(frame_izquierda, text="Volumen", padx=10, pady=5)
frame_volumen.grid(row=6, column=0, sticky="ew", padx=10, pady=(0, 10))
frame_volumen.columnconfigure(0, weight=1)

volumen_slider = tk.Scale(frame_volumen, from_=0, to=100, orient="horizontal", label="Nivel de Volumen")
volumen_slider.set(50)
volumen_slider.grid(row=0, column=0, sticky="ew")

# === Botones dinámicos por archivo === Creacioón de botón de archvios por carpeta
# frame_archivos = tk.LabelFrame(frame_izquierda, text="Archivos de Carpeta", padx=10, pady=5)
# frame_archivos.grid(row=8, column=0, sticky="ew", padx=10, pady=(0, 10))

# btn_cargar_carpeta = tk.Button(frame_archivos, text="Seleccionar carpeta", command=crear_botones_desde_carpeta)
# btn_cargar_carpeta.pack(fill="x", padx=5, pady=5)

''' DANI
# === Botones de videos encontrados en carpeta ===
frame_videos_dinamicos = tk.LabelFrame(frame_izquierda, text="Videos de Carpeta", padx=10, pady=5)
frame_videos_dinamicos.grid(row=9, column=0, sticky="ew", padx=10, pady=(0, 10))

btn_cargar_videos = tk.Button(
    frame_videos_dinamicos,
    text="Seleccionar carpeta de videos",
    command=crear_botones_videos_desde_carpeta
)
btn_cargar_videos.pack(fill="x", padx=5, pady=5)

frame_galeria_videos = tk.Frame(frame_videos_dinamicos)
frame_galeria_videos.pack(fill="both", expand=True)
'''



##MARTA
# === Botones de videos neutros en carpeta ===
frame_videos_dinamicos = tk.LabelFrame(frame_izquierda, text="Videos 360 - Escenarios Neutros", padx=10, pady=5)
frame_videos_dinamicos.grid(row=9, column=0, sticky="ew", padx=10, pady=(0, 10))

# Subframe para los botones de cargar y borrar en la misma fila pero en dos columnas
frame_botones = tk.Frame(frame_videos_dinamicos)
frame_botones.pack(fill="x", padx=5, pady=5)

def subir_videos_ansiosos():
    archivos = filedialog.askopenfilenames(
        title="Selecciona videos ansiosos para subir a las gafas",
        filetypes=[("Archivos de video", "*.mp4 *.avi *.mov *.mkv")]
    )

    if not archivos:
        print("No se seleccionaron archivos.")
        return

    limpiar_directorio_temporal()

    ventana_progreso = tk.Toplevel()
    ventana_progreso.title("Subiendo videos ansiosos")
    ventana_progreso.geometry("400x100")
    ventana_progreso.resizable(False, False)

    etiqueta_estado = tk.Label(ventana_progreso, text="Preparando...", anchor="w")
    etiqueta_estado.pack(fill="x", padx=10, pady=(10, 0))

    barra = ttk.Progressbar(ventana_progreso, maximum=len(archivos), mode="determinate")
    barra.pack(fill="x", padx=10, pady=10)

    def hilo_subida():
        ruta_remota = "/sdcard/Movies/EscenariosAnsiosos"  # CAMBIADA

        for i, ruta_local in enumerate(archivos):
            nombre_original = os.path.basename(ruta_local)
            nombre_normalizado = normalizar_nombre(nombre_original)

            if nombre_normalizado != nombre_original:
                ruta_temp = os.path.join("tmp", nombre_normalizado)
                os.makedirs("tmp", exist_ok=True)
                shutil.copy(ruta_local, ruta_temp)
                ruta_local_usar = ruta_temp
            else:
                ruta_local_usar = ruta_local

            ruta_remota_video = f"{ruta_remota}/{nombre_normalizado}"

            def actualizar_estado(msg, progreso=None):
                def _actualizar():
                    etiqueta_estado.config(text=msg)
                    if progreso is not None:
                        barra['value'] = progreso
                ventana_progreso.after(0, _actualizar)

            try:
                actualizar_estado(f"Subiendo: {nombre_normalizado}", i)
                adb_push(ruta_local_usar, ruta_remota_video)

                miniatura = generar_miniatura_local(ruta_local)
                if miniatura:
                    nombre_thumb = os.path.splitext(nombre_normalizado)[0] + ".jpg"
                    ruta_thumb_local = os.path.join("tmp", nombre_thumb)
                    miniatura.save(ruta_thumb_local)
                    ruta_remota_thumb = f"{ruta_remota}/{nombre_thumb}"
                    adb_push(ruta_thumb_local, ruta_remota_thumb)

                actualizar_estado(f"Subido: {nombre_normalizado}", i + 1)

            except Exception as e:
                print(f"Error al subir {ruta_local}: {e}")
                actualizar_estado(f"Error con {nombre_normalizado}", i + 1)

        ventana_progreso.after(0, lambda: etiqueta_estado.config(text="Subida completada."))
        ventana_progreso.after(1000, ventana_progreso.destroy)

    threading.Thread(target=hilo_subida, daemon=True).start()


def subir_videos_ansiosos():
    subir_videos_gafas(tipo="ansioso")

def crear_botones_videos_ansiosos_en_frame():
    videos = cargar_videos_ansiosos()

    if not videos:
        print("No se encontraron videos ansiosos. Saliendo de la función.")
        return

    for widget in frame_galeria_videos_ansiosos.winfo_children():
        widget.destroy()

    imagenes_miniaturas.clear()
    columnas = 4
    fila = 0
    columna = 0

    font_texto = tkFont.Font(family="Arial", size=10)

    img_transparente = Image.new('RGBA', (160, 90), (255, 255, 255, 0))
    img_transparente_tk = ImageTk.PhotoImage(img_transparente)
    imagenes_miniaturas.append(img_transparente_tk)

    for video in videos:
        base_name = video.rsplit('.', 1)[0]
        ruta_img_jpg = f"/sdcard/Movies/EscenariosAnsiosos/{base_name}.jpg"
        ruta_img_png = f"/sdcard/Movies/EscenariosAnsiosos/{base_name}.png"

        ruta_imagen = None
        if archivo_existe_en_dispositivo(ruta_img_jpg):
            ruta_imagen = copiar_imagen_desde_dispositivo(ruta_img_jpg)
        elif archivo_existe_en_dispositivo(ruta_img_png):
            ruta_imagen = copiar_imagen_desde_dispositivo(ruta_img_png)

        marco_video = tk.Frame(frame_galeria_videos_ansiosos)

        if ruta_imagen:
            try:
                img = Image.open(ruta_imagen)
                img.thumbnail((160, 90))
                thumbnail = ImageTk.PhotoImage(img)
                imagenes_miniaturas.append(thumbnail)

                btn = tk.Button(marco_video, image=thumbnail,
                                command=lambda path=f"/sdcard/Movies/EscenariosAnsiosos/{video}": abrir_video(path))
                btn.image = thumbnail
                btn.pack()
            except Exception as e:
                print("Error mostrando imagen ansiosa:", e)
        else:
            btn = tk.Button(marco_video, image=img_transparente_tk,
                            command=lambda path=f"/sdcard/Movies/EscenariosAnsiosos/{video}": abrir_video(path))
            btn.image = img_transparente_tk
            btn.pack()

        etiqueta = tk.Label(marco_video, text=base_name, font=font_texto, wraplength=160)
        etiqueta.pack()
        marco_video.grid(row=fila, column=columna, padx=5, pady=5)

        columna += 1
        if columna >= columnas:
            columna = 0
            fila += 1




btn_cargar_videos = tk.Button(
    frame_botones,
    text="Subir videos",
    command=subir_videos_ansiosos
)


btn_borrar_videos = tk.Button(
    frame_botones,
    text="Borrar vídeos",
    command=pestaña_videos_gafas
)

btn_cargar_videos.grid(row=0, column=0, padx=5, pady=5, sticky="ew")
btn_borrar_videos.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

frame_botones.grid_columnconfigure(0, weight=1)
frame_botones.grid_columnconfigure(1, weight=1)

btn_cargar_videos = tk.Button(
    frame_videos_dinamicos,
    text="Cargar vídeos en la interfaz",
    command=crear_botones_videos_desde_carpeta
)
btn_cargar_videos.pack(fill="x", padx=5, pady=5)

frame_galeria_videos = tk.Frame(frame_videos_dinamicos)
frame_galeria_videos.pack(fill="both", expand=True)

# === Botones de videos ansiosos en carpeta ===
frame_videos_dinamicos_ansioso = tk.LabelFrame(frame_izquierda, text="Videos 360 - Escenarios Ansiosos", padx=10, pady=5)
frame_videos_dinamicos_ansioso.grid(row=10, column=0, sticky="ew", padx=10, pady=(0, 10))

# Subframe para los botones de cargar y borrar en la misma fila pero en dos columnas
frame_botones = tk.Frame(frame_videos_dinamicos_ansioso)
frame_botones.pack(fill="x", padx=5, pady=5)

btn_cargar_videos = tk.Button(
    frame_botones,
    text="Subir videos",
    command=subir_videos_ansiosos
)

btn_borrar_videos = tk.Button(
    frame_botones,
    text="Borrar vídeos",
    command=pestaña_videos_gafas
)

btn_cargar_videos.grid(row=0, column=0, padx=5, pady=5, sticky="ew")
btn_borrar_videos.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

frame_botones.grid_columnconfigure(0, weight=1)
frame_botones.grid_columnconfigure(1, weight=1)


#ANSIOSO
btn_cargar_videos = tk.Button(
    frame_videos_dinamicos_ansioso,
    text="Cargar vídeos en la interfaz",
    command=crear_botones_videos_ansiosos_en_frame
)
btn_cargar_videos.pack(fill="x", padx=5, pady=5)

frame_galeria_videos = tk.Frame(frame_videos_dinamicos)
frame_galeria_videos.pack(fill="both", expand=True)



frame_galeria_videos_ansiosos = tk.Frame(frame_videos_dinamicos_ansioso)
frame_galeria_videos_ansiosos.pack(fill="both", expand=True)
'''
# === Botones de imágenes encontradas en carpeta ===
frame_imagenes_dinamicas = tk.LabelFrame(frame_izquierda, text="Imágenes de Carpeta", padx=10, pady=5)
frame_imagenes_dinamicas.grid(row=11, column=0, sticky="ew", padx=10, pady=(0, 10))

btn_cargar_imagenes = tk.Button(frame_imagenes_dinamicas, text="Seleccionar carpeta de imágenes",
                                command=crear_botones_imagenes_desde_carpeta)
btn_cargar_imagenes.pack(fill="x", padx=5, pady=5)

imagen_preview_label = tk.Label(frame_imagenes_dinamicas, bg="white", relief="sunken")
imagen_preview_label.pack(fill="both", padx=5, pady=5)
'''

# === Votacion ===
frame_votacion = tk.Frame(frame_izquierda)
frame_votacion.grid(row=8, column=0, sticky="ew", padx=10, pady=(0, 10))
frame_votacion.columnconfigure((0, 1, 2), weight=1)

tk.Button(frame_votacion, text="Lanzar votación", command=lambda: registrar_boton("Lanzar votación")).grid(row=0,
                                                                                                           column=0,
                                                                                                           padx=5,
                                                                                                           pady=2,
                                                                                                           sticky="ew")
tk.Button(frame_votacion, text="Parar Votación", command=lambda: registrar_boton("Parar votación")).grid(row=0,
                                                                                                         column=1,
                                                                                                         padx=5, pady=2,
                                                                                                         sticky="ew")
tk.Button(frame_votacion, text="Votación errónea", command=lambda: registrar_boton("Votación errónea")).grid(row=0,
                                                                                                             column=2,
                                                                                                             padx=5,
                                                                                                             pady=2,
                                                                                                             sticky="ew")
'''
# === Gafas (parte derecha) ===
frame_gafas = tk.LabelFrame(frame_contenido, text="Vista de Gafas", bg="white", padx=5, pady=5)
frame_gafas.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)
frame_gafas.columnconfigure(0, weight=1)

try:
    imagen_gafas_original = Image.open("gafas.png")
except FileNotFoundError:
    imagen_gafas_original = Image.new("RGB", (400, 300), "gray")
    d = ImageDraw.Draw(imagen_gafas_original)
    d.text((100, 130), "Sin imagen", fill="white")

label_gafas = tk.Label(frame_gafas, bg="white", relief="solid", bd=1)
label_gafas.grid(row=0, column=0, sticky="nsew")
'''
'''
def redimensionar_gafas(event):
    max_alto = 400
    ancho = frame_gafas.winfo_width()
    alto_original = imagen_gafas_original.height
    ancho_original = imagen_gafas_original.width
    escala = min(ancho / ancho_original, max_alto / alto_original)
    nuevo_ancho = int(ancho_original * escala)
    nuevo_alto = int(alto_original * escala)
    imagen_redimensionada = imagen_gafas_original.resize((nuevo_ancho, nuevo_alto), Image.LANCZOS)
    imagen_tk = ImageTk.PhotoImage(imagen_redimensionada)
    label_gafas.config(image=imagen_tk)
    label_gafas.image = imagen_tk
'''

def actualizar_sesion_en_formulario():
    if participante_actual:
        entry_id.delete(0, tk.END)
        entry_id.insert(0, participante_actual)

        sesion = participantes.get(participante_actual, 1)
        entry_sesion.configure(state="normal")
        entry_sesion.delete(0, tk.END)
        entry_sesion.insert(0, str(sesion))
        entry_sesion.configure(state="readonly")


# === Generar Informe ===
btn_generar_informe = tk.Button(frame_izquierda, text="Generar informe", bg="#4da6ff", fg="white",
                                font=("Segoe UI", 10, "bold"), command=mostrar_informe)
btn_generar_informe.grid(row=11, column=0, sticky="ew", padx=10, pady=(0, 10))
frame_izquierda.rowconfigure(11, weight=1)

from tkinter import simpledialog


def ver_vista_tablet():
    import cv2
    from PIL import Image, ImageTk
    import numpy as np
    from tkinter import Toplevel, Label

    # Obtener la IP automáticamente desde el primer dispositivo ADB conectado
    devices = list_devices_tcpip()
    if not devices:
        messagebox.showerror("Sin conexión", "No se detectó ninguna tablet conectada por ADB en red.")
        return

    ip = get_device_ip(devices[0])  # Ya tienes esta función en tu código
    if not ip:
        messagebox.showerror("Error", "No se pudo obtener la IP de la tablet.")
        return

    url = f"http://{ip}:8080/stream.mjpeg"  # URL específica de ScreenStream

    cap = cv2.VideoCapture(url)
    if not cap.isOpened():
        messagebox.showerror("Error", f"No se pudo abrir el stream desde {url}")
        return

    ventana = Toplevel()
    ventana.title(f"Vista de la Tablet ({ip})")
    lbl = Label(ventana)
    lbl.pack()

    def actualizar():
        if cap.isOpened():
            ret, frame = cap.read()
            if ret:
                frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                img = Image.fromarray(frame).resize((800, 480))
                imgtk = ImageTk.PhotoImage(image=img)
                lbl.imgtk = imgtk
                lbl.configure(image=imgtk)
        ventana.after(30, actualizar)

    actualizar()

    def cerrar():
        cap.release()
        ventana.destroy()

    ventana.protocol("WM_DELETE_WINDOW", cerrar)


# TABLET
btn_vista_tablet = tk.Button(frame_dispositivo, text="Ver vista de los biosensores", command=ver_vista_tablet)
btn_vista_tablet.grid(row=4, column=0, columnspan=2, pady=5, sticky="ew")

combo_participantes['values'] = actualizar_lista_participantes()
'''
frame_gafas.bind("<Configure>", redimensionar_gafas)
'''
root.bind("<Configure>", reorganizar_botones)
actualizar_timer()

root.mainloop()
